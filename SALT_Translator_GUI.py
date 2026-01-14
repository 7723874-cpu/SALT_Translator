import os
import threading

from tkinter import (
    Frame,
    Label,
    Button,
    END,
    DISABLED,
    NORMAL,
    BOTH,
    LEFT,
    X,
    BOTTOM,
    messagebox,
    StringVar,
)
from tkinter.scrolledtext import ScrolledText
from tkinter import filedialog

from dotenv import load_dotenv
from openai import OpenAI
from docx import Document
from docx.shared import Pt

# Попытка подключить drag & drop
try:
    from tkinterdnd2 import DND_FILES, TkinterDnD
    DND_AVAILABLE = True
except ImportError:
    from tkinter import Tk  # обычный Tk, без dnd
    DND_AVAILABLE = False

# =======================
#  Инициализация OpenAI
# =======================
load_dotenv()
API_KEY = os.getenv("OPENAI_API_KEY")

if not API_KEY:
    print("ВНИМАНИЕ: Переменная окружения OPENAI_API_KEY не найдена.")

client = OpenAI(api_key=API_KEY)

BASE_SYSTEM_PROMPT = """
You are a professional bilingual translator (English <-> Russian) with accuracy equal to or better than DeepL.

General rules:
- Preserve meaning, nuance, tone and style.
- Use natural, conversational, flowing language with rather short sentences.
- Do NOT output the original text, only the final translation.
- No comments, explanations, or labels — translation only.
"""


def build_system_prompt(mode: str) -> str:
    """
    mode: 'auto' | 'en-ru' | 'ru-en'
    """
    if mode == "en-ru":
        extra = "Translate all user text strictly from English to Russian."
    elif mode == "ru-en":
        extra = "Translate all user text strictly from Russian to English."
    else:
        extra = (
            "If the source text is in English, translate it into Russian. "
            "If the source text is in Russian, translate it into English."
        )
    return BASE_SYSTEM_PROMPT + "\n\n" + extra


def translate_text(text: str, mode: str) -> str:
    """
    Перевод через OpenAI (EN <-> RU) с учётом выбранного направления.
    mode: 'auto' | 'en-ru' | 'ru-en'
    """
    if not API_KEY:
        raise RuntimeError(
            "OPENAI_API_KEY не найден. Добавь ключ в .env или переменные среды."
        )

    if not text.strip():
        return ""

    system_prompt = build_system_prompt(mode)

    response = client.chat.completions.create(
        model="gpt-4.1",   # можно сменить на "gpt-4.1-mini" для экономии
        temperature=0.0,
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": text},
        ],
    )
    return response.choices[0].message.content.strip()


def read_file_text(path: str) -> str:
    """
    Читаем исходник как текст.
    Поддержка .docx и обычных текстовых (.txt, .srt, .md, .log, .csv и т.д.).
    """
    ext = os.path.splitext(path)[1].lower()

    if ext == ".docx":
        doc = Document(path)
        return "\n".join(p.text for p in doc.paragraphs)
    else:
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()


def save_translation_to_file(src_path: str, translated: str, mode: str) -> str:
    """
    Сохранение перевода рядом с исходным файлом.
    Имя вида: name_ru.ext / name_en.ext / name_tr.ext.

    Для .docx:
    - одиночный интервал (1.0),
    - без отступов до/после абзаца,
    - без лишних пустых строк.
    """
    folder = os.path.dirname(src_path)
    base_name = os.path.splitext(os.path.basename(src_path))[0]
    ext = os.path.splitext(src_path)[1].lower()

    if mode == "en-ru":
        suffix = "_ru"
    elif mode == "ru-en":
        suffix = "_en"
    else:
        suffix = "_tr"

    out_name = f"{base_name}{suffix}{ext if ext else '.txt'}"
    out_path = os.path.join(folder, out_name)

    if ext == ".docx":
        doc = Document()

        normal_style = doc.styles["Normal"]
        pf = normal_style.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing = 1.0  # single

        # Если захочешь — можно включить фиксированный шрифт:
        # font = normal_style.font
        # font.name = "Times New Roman"
        # font.size = Pt(12)

        for line in translated.splitlines():
            if not line.strip():
                continue
            doc.add_paragraph(line.strip(), style="Normal")

        doc.save(out_path)
    else:
        with open(out_path, "w", encoding="utf-8", errors="ignore") as f:
            f.write(translated)

    return out_path


# =======================
#       GUI КЛИЕНТ
# =======================

class BatchTranslatorGUI:
    """
    Батч-переводчик:
    - Drag & Drop нескольких файлов
    - Кнопка "Выбрать файлы..."
    - Кнопка "Запустить перевод" руками
    - Никакого содержимого, только лог
    """

    def __init__(self, master):
        self.master = master
        master.title("SALT Batch Translator (OpenAI)")
        master.geometry("800x500")

        # режим перевода: auto / en-ru / ru-en
        self.mode_var = StringVar(value="auto")

        # очередь файлов на обработку
        self.queue = []
        self.is_processing = False

        # -------- Верхняя панель --------
        top_frame = Frame(master)
        top_frame.pack(fill=X, padx=10, pady=(10, 5))

        title_label = Label(
            top_frame,
            text="SALT Batch Translator — EN ↔ RU (OpenAI)",
            font=("Segoe UI", 14, "bold")
        )
        title_label.pack(side=LEFT)

        # -------- Панель выбора направления --------
        mode_frame = Frame(master)
        mode_frame.pack(fill=X, padx=10, pady=(0, 5))

        mode_title = Label(
            mode_frame,
            text="Направление перевода:",
            font=("Segoe UI", 9, "bold")
        )
        mode_title.pack(side=LEFT)

        from tkinter import Radiobutton

        Radiobutton(
            mode_frame,
            text="Auto (EN ↔ RU)",
            variable=self.mode_var,
            value="auto"
        ).pack(side=LEFT, padx=(10, 0))

        Radiobutton(
            mode_frame,
            text="EN → RU",
            variable=self.mode_var,
            value="en-ru"
        ).pack(side=LEFT, padx=(10, 0))

        Radiobutton(
            mode_frame,
            text="RU → EN",
            variable=self.mode_var,
            value="ru-en"
        ).pack(side=LEFT, padx=(10, 0))

        # -------- Зона дропа / описание --------
        drop_label = Label(
            master,
            text=(
                "Перетащи сюда один или несколько файлов,\n"
                "или нажми «Выбрать файлы…» ниже.\n"
                "Перевод будет сохранён рядом с каждым исходным файлом."
            ),
            font=("Segoe UI", 10),
            fg="gray"
        )
        drop_label.pack(fill=X, padx=10, pady=(5, 5))

        # -------- Лог --------
        self.log_text = ScrolledText(master, wrap="word", height=15, font=("Segoe UI", 9))
        self.log_text.pack(fill=BOTH, expand=True, padx=10, pady=(0, 5))
        self.log_text.config(state=DISABLED)

        # Drag & Drop на лог (зона дропа)
        if DND_AVAILABLE:
            try:
                self.log_text.drop_target_register(DND_FILES)
                self.log_text.dnd_bind("<<Drop>>", self.on_files_drop)
            except Exception as e:
                print("Не удалось инициализировать drag & drop:", e)

        # -------- Нижняя панель с кнопками --------
        bottom_frame = Frame(master)
        bottom_frame.pack(fill=X, padx=10, pady=(0, 5))

        select_button = Button(
            bottom_frame,
            text="Выбрать файлы…",
            command=self.on_select_files_click,
            width=16
        )
        select_button.pack(side=LEFT)

        self.process_button = Button(
            bottom_frame,
            text="Запустить перевод",
            command=self.on_process_click,
            width=16,
            state=NORMAL
        )
        self.process_button.pack(side=LEFT, padx=(10, 0))

        clear_log_button = Button(
            bottom_frame,
            text="Очистить лог",
            command=self.clear_log,
            width=12
        )
        clear_log_button.pack(side=LEFT, padx=(10, 0))

        # -------- Статус --------
        self.status_label = Label(
            master,
            text="Готово.",
            anchor="w",
            font=("Segoe UI", 9)
        )
        self.status_label.pack(fill=X, side=BOTTOM, padx=10, pady=(0, 5))

    # =======================
    #        ЛОГ
    # =======================

    def log(self, text: str):
        self.log_text.config(state=NORMAL)
        self.log_text.insert(END, text + "\n")
        self.log_text.see(END)
        self.log_text.config(state=DISABLED)

    def clear_log(self):
        self.log_text.config(state=NORMAL)
        self.log_text.delete("1.0", END)
        self.log_text.config(state=DISABLED)

    # =======================
    #    Обработка файлов
    # =======================

    def on_select_files_click(self):
        paths = filedialog.askopenfilenames(
            title="Выберите файлы для перевода",
            filetypes=(
                ("Текстовые файлы", "*.txt *.srt *.md *.log *.csv"),
                ("Документы Word", "*.docx"),
                ("Все файлы", "*.*"),
            )
        )
        if paths:
            self.add_files_to_queue(list(paths))

    def on_files_drop(self, event):
        """
        Drag & Drop нескольких файлов.
        event.data, как правило: "{C:/path/file 1.docx} {C:/path/file 2.docx}"
        """
        raw = event.data.strip()
        files = self._split_dnd_paths(raw)
        self.add_files_to_queue(files)

    @staticmethod
    def _split_dnd_paths(data: str):
        """
        Разбор строки от tkdnd в список путей.
        """
        result = []
        current = ""
        in_braces = False

        for ch in data:
            if ch == "{":
                in_braces = True
                current = ""
            elif ch == "}":
                in_braces = False
                if current:
                    result.append(current)
                    current = ""
            elif ch == " " and not in_braces:
                if current:
                    result.append(current)
                    current = ""
            else:
                current += ch

        if current:
            result.append(current)

        # фильтруем только реальные файлы
        result = [p for p in result if os.path.isfile(p)]
        return result

    def add_files_to_queue(self, files):
        """
        Добавляет файлы в очередь, НО НЕ запускает перевод.
        Перевод только по кнопке "Запустить перевод".
        """
        if not files:
            return

        added = 0
        for path in files:
            if path not in self.queue:
                self.queue.append(path)
                self.log(f"Добавлен в очередь: {path}")
                added += 1

        if added == 0:
            self.log("Все выбранные файлы уже были в очереди.")

        self.status_label.config(
            text=f"Файлов в очереди: {len(self.queue)}. Нажми «Запустить перевод»."
        )

    def on_process_click(self):
        if self.is_processing:
            messagebox.showinfo("Перевод", "Перевод уже выполняется.")
            return

        if not self.queue:
            messagebox.showinfo("Перевод", "Очередь пуста. Добавь файлы для перевода.")
            return

        self.start_processing()

    def start_processing(self):
        if not self.queue:
            return
        self.is_processing = True
        self.process_button.config(state=DISABLED)
        self.status_label.config(text="Перевод выполняется…")

        thread = threading.Thread(target=self._process_queue_in_thread)
        thread.daemon = True
        thread.start()

    def _process_queue_in_thread(self):
        mode = self.mode_var.get()
        while self.queue:
            path = self.queue.pop(0)
            try:
                self._process_single_file(path, mode)
            except Exception as e:
                self.master.after(
                    0,
                    self.log,
                    f"[ОШИБКА] {os.path.basename(path)}: {e}"
                )

        self.master.after(0, self._processing_finished)

    def _process_single_file(self, path: str, mode: str):
        filename = os.path.basename(path)
        self.master.after(0, self.log, f"Обработка: {filename}")

        text = read_file_text(path)
        translated = translate_text(text, mode)
        out_path = save_translation_to_file(path, translated, mode)

        self.master.after(
            0,
            self.log,
            f"✅ Готово: {filename} → {os.path.basename(out_path)}"
        )

    def _processing_finished(self):
        self.is_processing = False
        self.process_button.config(state=NORMAL)
        self.status_label.config(text="Готово. Очередь пуста.")

    # =======================


def main():
    if DND_AVAILABLE:
        root = TkinterDnD.Tk()
    else:
        root = Tk()
        print(
            "Drag & Drop недоступен (модуль tkinterdnd2 не установлен). "
            "Установи: pip install tkinterdnd2"
        )

    app = BatchTranslatorGUI(root)
    root.mainloop()


if __name__ == "__main__":
    main()
